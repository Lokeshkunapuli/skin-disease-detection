import json
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn


PROJECT_NAME = "Skin Disease Detection at the Edge Using OpenVINO"
OUTPUT_FILE = "research.docx"


def add_heading(document: Document, text: str, level: int = 1):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    return p


def add_body(document: Document, text: str):
    p = document.add_paragraph(text)
    p_format = p.paragraph_format
    p_format.space_after = Pt(6)
    return p


def add_figure_caption(document: Document, caption: str):
    p = document.add_paragraph()
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_reference(paragraph, index: int, text: str):
    paragraph.add_run(f"[{index}] ")
    paragraph.add_run(text)


def read_classes():
    classes_path = os.path.join("model_knn", "classes.json")
    if os.path.exists(classes_path):
        try:
            with open(classes_path, "r", encoding="utf-8") as f:
                names = json.load(f)
            if isinstance(names, list):
                return names
        except Exception:
            return []
    return []


def build_document():
    doc = Document()

    # Basic page setup (approximate IEEE look & feel for word-processor; true two-column requires template)
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Title block
    title = doc.add_paragraph()
    run = title.add_run(PROJECT_NAME)
    run.bold = True
    run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    authors = doc.add_paragraph()
    authors.add_run("Author: Your Name").bold = True
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER

    affiliation = doc.add_paragraph("Affiliation: Your Institution, City, Country")
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_p = doc.add_paragraph(datetime.now().strftime("Date: %B %d, %Y"))
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # Abstract
    add_heading(doc, "Abstract", level=1)
    add_body(
        doc,
        (
            "This paper presents a practical system for skin disease classification that runs entirely on CPU "
            "using Intel's OpenVINO toolkit for efficient edge inference. The application ingests a dermatoscopic "
            "image, performs preprocessing, executes a lightweight convolutional network compiled to OpenVINO, and "
            "returns the predicted disease along with an interpretable occlusion-sensitivity heatmap and a clinical "
            "severity assessment. A classical K-Nearest Neighbors (KNN) fallback provides resilience when the IR model "
            "is missing or incompatible. The system emphasizes latency and deployability over dataset-level accuracy, "
            "achieving sub-second inference on commodity hardware."
        ),
    )

    # Keywords
    add_heading(doc, "Index Terms—", level=1)
    add_body(doc, "OpenVINO, edge AI, dermatology, skin lesion analysis, explainable AI, KNN.")

    # Introduction
    add_heading(doc, "I. Introduction", level=1)
    add_body(
        doc,
        (
            "Skin cancer and other dermatological conditions represent a major global health burden. Early triage and "
            "risk assessment can improve outcomes via timely specialist referral. Recent advances in deep learning have "
            "enabled computer-assisted diagnosis; however, compute and memory constraints often limit deployment at the "
            "edge. We implement an end-to-end system that performs low-latency inference on CPU via OpenVINO [1], "
            "deliverable as a Flask web application. The solution also integrates interpretable heatmaps based on "
            "occlusion sensitivity [2] and a clinically informed severity score to aid decision support."
        ),
    )

    # Related Work
    add_heading(doc, "II. Related Work", level=1)
    add_body(
        doc,
        (
            "OpenVINO has been widely adopted to optimize and deploy DNNs for CPU and heterogeneous devices [1]. "
            "Explainability in medical imaging often leverages saliency methods such as occlusion maps [2] and Grad-CAM [3]. "
            "Dermatology datasets such as ISIC facilitate benchmarking skin lesion classifiers [4]. Classical methods, "
            "including KNN with color-texture features, remain useful baselines and fallbacks [5]."
        ),
    )

    # System Overview
    add_heading(doc, "III. System Overview", level=1)
    add_body(
        doc,
        (
            "The system is implemented as a Flask application that exposes a simple web UI for image upload. On submit, the "
            "server reads the image and invokes the OpenVINO runtime to execute the IR model (XML+BIN). If model loading "
            "fails due to legacy IR versions, a KNN fallback computes per-class scores from color histograms and simple "
            "gradient features. For interpretability, an occlusion-sensitivity heatmap perturbs local regions to estimate "
            "their contribution to the predicted class. Additionally, a rule-based clinical severity score summarizes lesion "
            "characteristics (area ratio, border irregularity, color heterogeneity, asymmetry, edge density, texture complexity)."
        ),
    )

    # Methods
    add_heading(doc, "IV. Methods", level=1)
    add_heading(doc, "A. Data and Classes", level=2)
    classes = read_classes()
    if classes:
        add_body(doc, f"The deployment includes {len(classes)} classes, e.g., {', '.join(classes[:6])}{'...' if len(classes) > 6 else ''}.")
    else:
        add_body(doc, "The repository includes a sample data hierarchy under skin_data/ organized by disease class.")

    add_heading(doc, "B. Preprocessing", level=2)
    add_body(
        doc,
        (
            "Images are resized to the network's input resolution, channels are ordered to match the model, and a batch "
            "dimension is added prior to inference. For KNN, features include per-channel color histograms and a simple "
            "gradient-magnitude histogram with L2 normalization."
        ),
    )

    add_heading(doc, "C. OpenVINO Inference", level=2)
    add_body(
        doc,
        (
            "The Intel OpenVINO runtime reads the IR graph and compiles it for CPU execution. Inference outputs are converted "
            "to probabilities using a numerically stable softmax. The top-1 class and its probability (0–1) are reported as "
            "the predicted label and accuracy, respectively."
        ),
    )

    add_heading(doc, "D. KNN Fallback", level=2)
    add_body(
        doc,
        (
            "If the IR model cannot be loaded (e.g., legacy IR version), a KNN fallback computes per-class scores as "
            "inverse-distance accumulations over the k nearest neighbors. The scores are transformed with softmax to "
            "obtain a proper probability distribution, and the top-1 probability is presented as accuracy (0–1)."
        ),
    )

    add_heading(doc, "E. Explainability via Occlusion Sensitivity", level=2)
    add_body(
        doc,
        (
            "We adopt occlusion sensitivity [2]: sliding a patch over the image while re-scoring the model, quantifying how "
            "occluding a region lowers the target-class score. The resulting map is normalized to [0, 1] and overlaid on "
            "the image for transparency."
        ),
    )

    add_heading(doc, "F. Clinical Severity Scoring", level=2)
    add_body(
        doc,
        (
            "We define a lightweight severity score inspired by ABCD criteria, aggregating area ratio, border irregularity, "
            "color heterogeneity, asymmetry, edge density, and texture complexity into a 0–1 score with heuristic thresholds. "
            "The score is mapped to Low/Moderate/High/Very High levels accompanied by actionable advice."
        ),
    )

    # Experiments
    add_heading(doc, "V. Experiments", level=1)
    add_body(
        doc,
        (
            "Latency was measured on a commodity CPU using the OpenVINO path, yielding sub-second end-to-end response time "
            "(upload + preprocess + inference + rendering). As the repository prioritizes deployment, we report per-image "
            "confidence rather than dataset-level accuracy; the README explicitly notes the model is not tuned for high "
            "accuracy. Qualitative examples demonstrate that the occlusion heatmaps highlight salient lesion regions."
        ),
    )

    add_heading(doc, "VI. Discussion", level=1)
    add_body(
        doc,
        (
            "The system balances deployability, interpretability, and responsiveness. OpenVINO ensures efficient CPU inference, "
            "while the KNN fallback provides robustness. The clinical severity score and heatmaps can support triage, but "
            "they do not replace clinical judgment. Future work includes stronger backbones, better calibration, and external "
            "validation on diverse datasets."
        ),
    )

    add_heading(doc, "VII. Limitations and Ethical Considerations", level=1)
    add_body(
        doc,
        (
            "This system is intended for educational and assistive purposes, not as a diagnostic device. Dataset bias, image "
            "quality, and domain shift can affect predictions. Any deployment must ensure privacy, informed consent, and "
            "appropriate regulatory review."
        ),
    )

    add_heading(doc, "VIII. Conclusion", level=1)
    add_body(
        doc,
        (
            "We presented an edge-ready skin lesion analysis tool leveraging OpenVINO for fast CPU inference, with a classical "
            "KNN fallback, interpretable occlusion heatmaps, and a clinical severity score. The design demonstrates how to "
            "deliver responsive, explainable, and practical AI capabilities within constrained environments."
        ),
    )

    # References (IEEE style numbering)
    add_heading(doc, "References", level=1)
    refs = [
        "OpenVINO Toolkit Documentation, Intel, 2025. [Online]. Available: https://docs.openvino.ai/",
        "M. D. Zeiler and R. Fergus, 'Visualizing and Understanding Convolutional Networks,' in ECCV, 2014, pp. 818–833.",
        "R. R. Selvaraju et al., 'Grad-CAM: Visual Explanations from Deep Networks via Gradient-based Localization,' in ICCV, 2017, pp. 618–626.",
        "No Authors, 'ISIC Archive: International Skin Imaging Collaboration,' [Online]. Available: https://www.isic-archive.com/",
        "G. Bradski, 'The OpenCV Library,' Dr. Dobb’s Journal of Software Tools, 2000.",
    ]
    for i, r in enumerate(refs, start=1):
        p = doc.add_paragraph()
        add_reference(p, i, r)

    # Save
    doc.save(OUTPUT_FILE)


if __name__ == "__main__":
    build_document()


